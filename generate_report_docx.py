from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
RESULTS_DIR = ROOT / "results"
DATA_DIR = ROOT / "data" / "fiqa"
OUTPUT_PATH = ROOT / "BaiThucHanh7.docx"


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8") as file:
        return list(csv.DictReader(file))


def fmt(value: str | float | None) -> str:
    if value in (None, "", "-"):
        return "-" if value in (None, "-") else ""
    return f"{float(value):.4f}"


def add_paragraph(document: Document, text: str, *, bold: bool = False, italic: bool = False, align=None, space_after: int = 6):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph


def add_bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def set_cell_text(cell, text: str, *, bold: bool = False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)


def build_document() -> Document:
    assignment_rows = load_csv(RESULTS_DIR / "assignment_table.csv")
    metric_rows = load_csv(RESULTS_DIR / "metrics.csv")
    metadata = json.loads((DATA_DIR / "metadata.json").read_text(encoding="utf-8"))

    llm_rows = [row for row in assignment_rows if row["Method"] == "LLM-QLM"]
    hybrid_rows = [row for row in metric_rows if row["method"] == "Hybrid"]
    hybrid_rows.sort(key=lambda row: (int(float(row["k"] or 0)), float(row["alpha"] or 0.0)))
    best_overall = max(metric_rows, key=lambda row: float(row["ndcg@10"]))

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(13)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÀI THỰC HÀNH 7")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Re-ranking tài liệu bằng LLM-QLM trên bộ dữ liệu FiQA")
    run.bold = True
    run.font.size = Pt(15)

    add_paragraph(document, "Họ và tên: .................................................................")
    add_paragraph(document, "MSSV: ..........................................................................", space_after=10)

    add_paragraph(document, "1. Giới thiệu", bold=True, space_after=8)
    add_paragraph(
        document,
        "Trong hệ thống truy hồi thông tin, nhiệm vụ chính là tìm và xếp hạng các tài liệu liên quan đến truy vấn của người dùng. "
        "Các phương pháp truyền thống như BM25 thường dựa vào mức độ trùng khớp từ khóa giữa truy vấn và tài liệu. "
        "Cách tiếp cận này mạnh ở việc tìm kiếm nhanh và hiệu quả khi câu hỏi dùng đúng các từ xuất hiện trong tài liệu, "
        "nhưng có thể bỏ sót các tài liệu liên quan về ngữ nghĩa nếu cách diễn đạt khác nhau."
    )
    add_paragraph(
        document,
        "Bài thực hành này cài đặt lại ý tưởng từ bài báo “Open-source Large Language Models are Strong Zero-shot Query Likelihood Models for Document Ranking (EMNLP Findings 2023)”. "
        "Ý tưởng cốt lõi là dùng mô hình ngôn ngữ lớn để ước lượng xác suất truy vấn được sinh ra từ tài liệu theo tư tưởng Query Likelihood Model. "
        "Nếu một tài liệu khiến truy vấn có xác suất sinh ra cao, tài liệu đó được xem là phù hợp hơn với truy vấn."
    )
    add_paragraph(document, "Mục tiêu của bài thực hành gồm:")
    add_bullet(document, "Xây dựng baseline BM25 để làm mốc so sánh.")
    add_bullet(document, "Cài đặt LLM-QLM để re-rank tài liệu.")
    add_bullet(document, "So sánh BM25, LLM-QLM và Hybrid để đánh giá hiệu quả của việc kết hợp tín hiệu lexical và ngữ nghĩa.")

    add_paragraph(document, "2. Bộ dữ liệu (FiQA)", bold=True, space_after=8)
    add_paragraph(
        document,
        "FiQA là bộ dữ liệu truy hồi thông tin trong lĩnh vực tài chính. Bộ dữ liệu gồm các câu hỏi thực tế và các tài liệu liên quan, "
        "phù hợp để đánh giá các hệ thống retrieval vì truy vấn thường ngắn nhưng tài liệu có thể diễn đạt cùng ý theo nhiều cách khác nhau."
    )
    add_paragraph(document, "Số lượng dữ liệu thu được sau bước chuẩn bị:")
    add_bullet(document, f"Documents: {metadata['documents']}")
    add_bullet(document, f"Queries: {metadata['queries']}")
    add_bullet(document, f"Relevance judgments: {metadata['relevance_judgments']}")
    add_paragraph(document, "Ý nghĩa của các thành phần dữ liệu:")
    add_bullet(document, "Documents: tập văn bản dùng để truy hồi. Mỗi tài liệu được lưu với id, title, text và trường contents đã được chuẩn hóa.")
    add_bullet(document, "Queries: các câu hỏi tài chính mà hệ thống cần tìm tài liệu phù hợp.")
    add_bullet(document, "Relevance judgments (qrels): nhãn liên quan giữa query_id và doc_id, dùng làm ground truth để đánh giá.")
    add_paragraph(document, "Ví dụ:")
    add_bullet(document, "Query: What is a good strategy for long term investing?")
    add_bullet(document, "Document: Long term investors often diversify their portfolio across different asset classes.")
    add_bullet(document, "Relevance: Relevant")
    add_paragraph(
        document,
        "Sau khi chạy prepare_fiqa.py, hệ thống sinh ra các file corpus.jsonl, queries.tsv, qrels.tsv, qrels.trec và metadata.json trong thư mục data/fiqa. "
        "Đây là các file đầu vào cho các bước build index, truy hồi, re-ranking và đánh giá."
    )

    add_paragraph(document, "3. Phương pháp", bold=True, space_after=8)
    add_paragraph(document, "3.1 Baseline: BM25", bold=True, space_after=6)
    add_paragraph(
        document,
        "BM25 là mô hình lexical kinh điển trong retrieval. Mô hình này chấm điểm dựa trên sự xuất hiện của từ khóa truy vấn trong tài liệu, "
        "đồng thời hiệu chỉnh theo độ dài tài liệu để tránh ưu tiên quá mức các văn bản dài. "
        "Trong bài này, BM25 được dùng để truy hồi top-100 tài liệu đầu tiên cho mỗi truy vấn với tham số k1 = 0.9 và b = 0.4."
    )
    add_paragraph(
        document,
        "Luồng xử lý của baseline là: Query -> BM25 -> Top-k documents. "
        "Ưu điểm của BM25 là tốc độ nhanh và recall tốt khi từ khóa trùng nhau; nhược điểm là chưa phản ánh tốt quan hệ ngữ nghĩa giữa truy vấn và tài liệu."
    )

    add_paragraph(document, "3.2 LLM-QLM Re-ranking", bold=True, space_after=6)
    add_paragraph(
        document,
        "Sau khi có tập ứng viên từ BM25, hệ thống dùng LLM-QLM để sắp xếp lại các tài liệu này. "
        "Với mỗi tài liệu, prompt được tạo theo mẫu “Document: <nội dung tài liệu>\\nQuery:”. "
        "Mô hình ngôn ngữ sau đó tính tổng log-xác suất của các token trong truy vấn khi được sinh tiếp theo sau prompt, tức là score(q, d) = log P(q | d)."
    )
    add_paragraph(
        document,
        "Trong cài đặt hiện tại, mô hình dùng là distilgpt2 thông qua thư viện transformers. "
        "Để tránh vượt quá giới hạn ngữ cảnh, văn bản tài liệu được cắt tối đa 1200 ký tự và độ dài đầu vào tối đa là 1024 token. "
        "Việc chấm điểm được thực hiện theo batch size = 4 để giảm thời gian chạy."
    )
    add_paragraph(
        document,
        "Ý nghĩa trực quan của QLM là: tài liệu nào mô tả đúng bối cảnh của truy vấn hơn thì truy vấn sẽ có xác suất được mô hình “sinh ra” cao hơn. "
        "Nhờ vậy, LLM-QLM có thể cải thiện xếp hạng top đầu kể cả khi từ ngữ giữa truy vấn và tài liệu không trùng hoàn toàn."
    )

    add_paragraph(document, "3.3 Hybrid Ranking", bold=True, space_after=6)
    add_paragraph(
        document,
        "Hybrid Ranking kết hợp điểm BM25 và điểm QLM theo công thức: Score(q, d) = α × BM25_score + (1 − α) × QLM_score. "
        "Vì hai nguồn điểm này nằm trên các thang đo khác nhau, hệ thống chuẩn hóa điểm bằng min-max trước khi cộng."
    )
    add_paragraph(document, "Ý nghĩa của tham số α:")
    add_bullet(document, "α nhỏ: ưu tiên QLM nhiều hơn.")
    add_bullet(document, "α lớn: ưu tiên BM25 nhiều hơn.")
    add_bullet(document, "α trung bình: cân bằng giữa trùng khớp từ khóa và phù hợp ngữ nghĩa.")
    add_paragraph(document, "Trong bài này, các giá trị α được thử nghiệm là 0.2, 0.5 và 0.8.")

    add_paragraph(document, "4. Thiết lập thực nghiệm", bold=True, space_after=8)
    add_paragraph(
        document,
        "Ba nhóm phương pháp được so sánh là BM25, LLM-QLM và Hybrid. "
        "BM25 đóng vai trò truy hồi ban đầu, còn LLM-QLM và Hybrid chỉ re-rank trên top-k ứng viên mà BM25 đã lấy ra trước đó."
    )
    add_paragraph(document, "Các tham số thực nghiệm:")
    add_bullet(document, "Top-k documents dùng để re-rank: 10, 20, 50.")
    add_bullet(document, "Tham số α của Hybrid: 0.2, 0.5, 0.8.")
    add_bullet(document, "Model sử dụng: distilgpt2.")
    add_bullet(document, "Framework: transformers.")
    add_bullet(document, "Thiết bị chạy: CPU.")
    add_bullet(document, "BM25 baseline được lấy ở độ sâu 100 để tính Recall@100.")

    setup_table = document.add_table(rows=1, cols=2)
    setup_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    setup_table.style = "Table Grid"
    set_cell_text(setup_table.rows[0].cells[0], "Thông tin", bold=True)
    set_cell_text(setup_table.rows[0].cells[1], "Giá trị", bold=True)
    for left, right in [
        ("Model", "distilgpt2"),
        ("Size", "81,912,576 tham số"),
        ("Framework", "transformers"),
        ("Batch size", "4"),
        ("BM25 k1", "0.9"),
        ("BM25 b", "0.4"),
        ("Hybrid normalization", "min-max"),
    ]:
        row = setup_table.add_row().cells
        set_cell_text(row[0], left)
        set_cell_text(row[1], right)

    add_paragraph(document, "5. Chỉ số đánh giá", bold=True, space_after=8)
    add_paragraph(
        document,
        "Bài thực hành sử dụng hai chỉ số đánh giá chính là nDCG@10 và Recall@100."
    )
    add_bullet(document, "nDCG@10 đánh giá chất lượng xếp hạng trong top-10. Chỉ số này vừa xét đúng/sai, vừa xét vị trí của tài liệu liên quan trong danh sách.")
    add_bullet(document, "Recall@100 đo khả năng thu hồi tài liệu liên quan trong 100 tài liệu đầu. Chỉ số này phản ánh mức độ bao phủ của hệ thống.")
    add_paragraph(
        document,
        "Trong code, việc đánh giá được thực hiện bằng thư viện ranx. Nếu máy có trec_eval, chương trình cũng hỗ trợ đối chiếu kết quả theo chuẩn TREC."
    )

    add_paragraph(document, "6. Kết quả thực nghiệm", bold=True, space_after=8)
    add_paragraph(
        document,
        "Bảng dưới đây là bảng kết quả chính dùng để trình bày trong báo cáo, bám đúng form mà đề bài yêu cầu."
    )

    main_table = document.add_table(rows=1, cols=5)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_table.style = "Table Grid"
    for index, header in enumerate(("Method", "k", "α", "nDCG@10", "Recall@100")):
        set_cell_text(main_table.rows[0].cells[index], header, bold=True)
    for row_data in assignment_rows:
        row = main_table.add_row().cells
        set_cell_text(row[0], row_data["Method"])
        set_cell_text(row[1], row_data["k"])
        set_cell_text(row[2], row_data["alpha"])
        set_cell_text(row[3], fmt(row_data["nDCG@10"]))
        set_cell_text(row[4], fmt(row_data["Recall@100"]))

    add_paragraph(document, "Nhận xét trực tiếp từ bảng chính:")
    add_bullet(document, f"BM25 đạt nDCG@10 = {fmt(assignment_rows[0]['nDCG@10'])} và Recall@100 = {fmt(assignment_rows[0]['Recall@100'])}. Đây là baseline để so sánh.")
    add_bullet(document, f"LLM-QLM cải thiện dần khi tăng k: từ {fmt(llm_rows[0]['nDCG@10'])} ở k = 10 lên {fmt(llm_rows[2]['nDCG@10'])} ở k = 50.")
    add_bullet(document, "Trong ba cấu hình Hybrid đúng theo bảng đề bài, cấu hình tốt nhất là k = 20, α = 0.5 với nDCG@10 = 0.2710.")
    add_bullet(document, "Cấu hình Hybrid k = 50, α = 0.8 giảm chất lượng so với các cấu hình Hybrid khác, cho thấy nếu ưu tiên BM25 quá nhiều thì lợi ích của QLM bị giảm.")

    add_paragraph(
        document,
        "Ngoài bảng chính, file results/metrics.csv còn lưu toàn bộ các cấu hình Hybrid đã chạy. "
        "Bảng phụ dưới đây giúp giải thích sâu hơn các kết quả xuất ra để người đọc hiểu vì sao có sự khác biệt giữa các cấu hình."
    )

    hybrid_table = document.add_table(rows=1, cols=4)
    hybrid_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hybrid_table.style = "Table Grid"
    for index, header in enumerate(("k", "α", "nDCG@10", "Recall@100")):
        set_cell_text(hybrid_table.rows[0].cells[index], header, bold=True)
    for row_data in hybrid_rows:
        row = hybrid_table.add_row().cells
        set_cell_text(row[0], int(float(row_data["k"])))
        set_cell_text(row[1], fmt(row_data["alpha"]))
        set_cell_text(row[2], fmt(row_data["ndcg@10"]))
        set_cell_text(row[3], fmt(row_data["recall@100"]))

    add_paragraph(document, "Kết quả tốt nhất toàn bộ các run đã đánh giá trong metrics.csv là:")
    add_bullet(document, f"Phương pháp: {best_overall['method']}")
    add_bullet(document, f"k = {int(float(best_overall['k']))}, α = {fmt(best_overall['alpha'])}")
    add_bullet(document, f"nDCG@10 = {fmt(best_overall['ndcg@10'])}")
    add_bullet(document, f"Recall@100 = {fmt(best_overall['recall@100'])}")
    add_paragraph(
        document,
        "Điểm này không nằm trong bảng nộp chính của đề bài vì assignment_table.csv chỉ lấy các cấu hình đại diện theo đúng yêu cầu đề. "
        "Do đó cần hiểu rõ rằng assignment_table.csv là bảng để chép vào báo cáo, còn metrics.csv là bảng để phân tích sâu toàn bộ thí nghiệm."
    )

    add_paragraph(document, "Giải thích ý nghĩa các file kết quả xuất ra:")
    add_bullet(document, "results/assignment_table.csv: bảng kết quả rút gọn đúng form đề bài, dùng để trình bày trong báo cáo.")
    add_bullet(document, "results/metrics.csv: bảng tổng hợp đầy đủ tất cả các run đã đánh giá, hữu ích cho việc phân tích và so sánh sâu hơn.")
    add_bullet(document, "results/experiment_setup.json: lưu cấu hình thực nghiệm như dataset, tham số BM25, model QLM, top-k và alpha.")
    add_bullet(document, "results/evaluation_summary.json: file tóm tắt vị trí các artifact đánh giá chính.")
    add_bullet(document, "results/report_summary.md: bản tóm tắt kết quả tự động do script đánh giá sinh ra.")

    add_paragraph(document, "7. Phân tích kết quả", bold=True, space_after=8)
    add_paragraph(document, "7.1 LLM-QLM có cải thiện so với BM25 không?", bold=True, space_after=6)
    add_paragraph(
        document,
        "Có, nếu xét theo nDCG@10. BM25 đạt 0.2361, trong khi LLM-QLM tăng lên 0.2469 ở k = 10, 0.2542 ở k = 20 và 0.2553 ở k = 50. "
        "Điều này cho thấy bước re-ranking bằng LLM giúp đưa các tài liệu phù hợp hơn lên top đầu."
    )
    add_paragraph(
        document,
        "Tuy nhiên, Recall@100 của LLM-QLM thấp hơn BM25. Lý do là QLM chỉ re-rank trên top-k tài liệu ứng viên mà BM25 đã truy hồi, "
        "nên khi k nhỏ thì số lượng tài liệu liên quan có thể thu hồi cũng bị giới hạn. Vì vậy, cần nhìn đồng thời cả nDCG@10 và Recall@100 để đánh giá công bằng."
    )

    add_paragraph(document, "7.2 Giá trị k tốt nhất", bold=True, space_after=6)
    add_paragraph(
        document,
        "Nếu chỉ xét riêng LLM-QLM trong ba giá trị k = 10, 20, 50 thì k = 50 cho nDCG@10 cao nhất. "
        "Điều này hợp lý vì top-k lớn hơn cho phép mô hình có nhiều ứng viên hơn để lựa chọn lại."
    )
    add_paragraph(
        document,
        "Nếu xét toàn bộ các cấu hình Hybrid trong metrics.csv, cấu hình tốt nhất là Hybrid với k = 50 và α = 0.2. "
        "Điều này cho thấy việc mở rộng số ứng viên và để QLM đóng vai trò chính, nhưng vẫn giữ một phần tín hiệu BM25, là chiến lược hiệu quả nhất trong thực nghiệm này."
    )

    add_paragraph(document, "7.3 Hybrid có hiệu quả hơn không?", bold=True, space_after=6)
    add_paragraph(
        document,
        "Có. Các cấu hình Hybrid tốt thường đạt nDCG@10 cao hơn cả BM25 lẫn LLM-QLM đơn lẻ. "
        "Nguyên nhân là BM25 nắm bắt tốt sự trùng khớp từ khóa, còn QLM bổ sung khả năng đánh giá mức độ phù hợp về ngữ nghĩa. "
        "Khi kết hợp hợp lý, hai nguồn tín hiệu này bù đắp điểm yếu cho nhau."
    )
    add_paragraph(
        document,
        "Ngược lại, khi α quá lớn như 0.8, trọng số nghiêng quá mạnh về BM25 nên lợi ích của QLM giảm. "
        "Điều đó giải thích vì sao không phải mọi cấu hình Hybrid đều tốt hơn, mà hiệu quả còn phụ thuộc vào cách chọn α."
    )

    add_paragraph(document, "7.4 Ví dụ truy vấn minh họa", bold=True, space_after=6)
    add_paragraph(
        document,
        "Với truy vấn “How to diversify investments?”, BM25 có xu hướng ưu tiên các tài liệu chứa trực tiếp các từ như diversify, investments, portfolio. "
        "Trong khi đó, LLM-QLM còn có khả năng tăng điểm cho những tài liệu nói về phân bổ tài sản, giảm rủi ro, đầu tư dài hạn hoặc đa dạng hóa danh mục dù cách diễn đạt không hoàn toàn giống truy vấn. "
        "Hybrid là cách dung hòa hai góc nhìn này: vừa tận dụng sự chính xác của từ khóa, vừa tận dụng khả năng hiểu ngữ cảnh của mô hình ngôn ngữ."
    )

    add_paragraph(document, "8. Trace", bold=True, space_after=8)
    add_paragraph(
        document,
        "Các file trace lưu kết quả xếp hạng cuối cùng của từng phương pháp, đúng theo yêu cầu đề bài. "
        "Mỗi dòng trong file trace có 4 cột: query_id, doc_id, score, rank."
    )
    add_paragraph(document, "Các file trace bắt buộc gồm:")
    for trace_name in (
        "trace/bm25.txt",
        "trace/qlm_k10.txt",
        "trace/qlm_k20.txt",
        "trace/qlm_k50.txt",
        "trace/hybrid_k10_a02.txt",
        "trace/hybrid_k20_a05.txt",
        "trace/hybrid_k50_a08.txt",
    ):
        add_bullet(document, trace_name)
    add_paragraph(document, "Giải thích vai trò của từng nhóm file đầu ra:")
    add_bullet(document, "trace/*.txt: file rút gọn, dễ xem và đúng định dạng nộp bài.")
    add_bullet(document, "runs/*.json: lưu điểm chi tiết theo cấu trúc query -> doc -> score, thuận tiện cho xử lý bằng code.")
    add_bullet(document, "runs/*.trec: lưu theo chuẩn TREC để tương thích với trec_eval.")
    add_bullet(document, "results/*.csv và *.json: dùng để tổng hợp kết quả, lập bảng và giải thích trong báo cáo.")
    add_paragraph(
        document,
        "Ví dụ, khi mở trace/bm25.txt, ta thấy danh sách tài liệu được sắp xếp theo điểm BM25. "
        "Khi mở trace/qlm_k20.txt hoặc trace/hybrid_k20_a05.txt, ta sẽ thấy cùng truy vấn đó nhưng thứ tự tài liệu đã thay đổi do bước re-ranking. "
        "Đây là minh chứng trực tiếp cho tác dụng của LLM-QLM và Hybrid trong bài thực hành."
    )

    add_paragraph(document, "10. Tài liệu tham khảo", bold=True, space_after=8)
    add_paragraph(
        document,
        "Zhuang, S., Liu, B., Koopman, B., & Zuccon, G. (2023). Open-source Large Language Models are Strong Zero-shot Query Likelihood Models for Document Ranking. Findings of EMNLP."
    )
    add_paragraph(
        document,
        "Ghi chú: Báo cáo này đã được điền sẵn toàn bộ kết quả thực nghiệm. Khi nộp bài chỉ cần bổ sung Họ và tên và MSSV.",
        italic=True,
        space_after=0,
    )

    return document


def main():
    document = build_document()
    document.save(OUTPUT_PATH)
    print(f"Saved {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
